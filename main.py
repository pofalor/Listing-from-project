import copy
import json
import logging
import os
import re
import sys

import docx
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


LOG_FILE = "listing_from_project.log"
SETTINGS_FILE = "listing_settings.json"
LOGGER = logging.getLogger(__name__)
DEFAULT_SETTINGS = {
    "document": {
        "title": "Листинг проекта",
        "title_font_name": "Times New Roman",
        "title_font_size": 16,
        "title_color": "#000000",
        "title_bold": True,
        "info_font_name": "Times New Roman",
        "info_font_size": 11,
        "info_color": "#000000",
    },
    "layout": {
        "two_columns_spacing_inches": 0.25,
    },
    "one_column": {
        "file_name_font_name": "Times New Roman",
        "file_name_font_size": 14,
        "file_name_color": "#000000",
        "file_name_bold": True,
        "content_font_name": "Courier New",
        "content_font_size": 8,
        "content_color": "#000000",
        "content_line_spacing": 1,
        "content_space_after": 0,
    },
    "two_columns": {
        "file_name_font_name": "Times New Roman",
        "file_name_font_size": 10,
        "file_name_color": "#000000",
        "file_name_bold": True,
        "content_font_name": "Courier New",
        "content_font_size": 6,
        "content_color": "#000000",
        "content_line_spacing": 1,
        "content_space_after": 0,
    },
}


def configure_logging(log_file=LOG_FILE):
    if LOGGER.handlers:
        return

    LOGGER.setLevel(logging.INFO)
    formatter = logging.Formatter("%(asctime)s [%(levelname)s] %(message)s")

    file_handler = logging.FileHandler(log_file, encoding='utf-8-sig')
    file_handler.setFormatter(formatter)
    LOGGER.addHandler(file_handler)

    stream_handler = logging.StreamHandler(sys.stdout)
    stream_handler.setFormatter(formatter)
    LOGGER.addHandler(stream_handler)

    LOGGER.propagate = False


def inches_to_twips(value):
    return int(value * 1440)


def set_section_columns(section, columns=1, spacing_inches=0.25):
    cols = section._sectPr.find(qn('w:cols'))
    if cols is None:
        cols = OxmlElement('w:cols')
        section._sectPr.append(cols)

    cols.set(qn('w:num'), str(columns))
    cols.set(qn('w:equalWidth'), '1')
    if columns > 1:
        cols.set(qn('w:space'), str(inches_to_twips(spacing_inches)))


def set_run_font(run, font_name, size=None):
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:ascii'), font_name)
    run._element.rPr.rFonts.set(qn('w:hAnsi'), font_name)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size is not None:
        run.font.size = Pt(size)


def deep_merge_settings(defaults, overrides):
    settings = copy.deepcopy(defaults)
    for key, value in overrides.items():
        if isinstance(value, dict) and isinstance(settings.get(key), dict):
            settings[key] = deep_merge_settings(settings[key], value)
        else:
            settings[key] = value
    return settings


def load_settings(settings_file_path=SETTINGS_FILE):
    if not os.path.exists(settings_file_path):
        LOGGER.warning("Файл настроек %s не найден. Используем настройки по умолчанию", settings_file_path)
        return copy.deepcopy(DEFAULT_SETTINGS)

    try:
        with open(settings_file_path, 'r', encoding='utf-8') as f:
            user_settings = json.load(f)
        settings = deep_merge_settings(DEFAULT_SETTINGS, user_settings)
        LOGGER.info("Настройки загружены из %s", settings_file_path)
        return settings
    except Exception as e:
        LOGGER.exception("Ошибка чтения файла настроек %s: %s", settings_file_path, e)
        LOGGER.warning("Используем настройки по умолчанию")
        return copy.deepcopy(DEFAULT_SETTINGS)


def parse_color(color_value):
    if not color_value:
        return None

    value = str(color_value).strip().lstrip('#')
    if not re.fullmatch(r'[0-9a-fA-F]{6}', value):
        LOGGER.warning("Некорректный цвет '%s'. Ожидается формат #RRGGBB", color_value)
        return None

    return RGBColor(
        int(value[0:2], 16),
        int(value[2:4], 16),
        int(value[4:6], 16),
    )


def apply_run_style(run, font_name, font_size=None, color=None, bold=None):
    if font_name:
        set_run_font(run, font_name, font_size)
    elif font_size is not None:
        run.font.size = Pt(font_size)

    rgb_color = parse_color(color)
    if rgb_color is not None:
        run.font.color.rgb = rgb_color
    if bold is not None:
        run.bold = bold


def load_ignore_patterns(ignore_file_path='.docignore'):
    # Загружает паттерны игнорирования из файла
    ignore_patterns = []
    if not os.path.exists(ignore_file_path):
        raise FileNotFoundError(
            f"Файл {ignore_file_path} не найден. Создайте файл .docignore для настройки игнорирования")

    try:
        with open(ignore_file_path, 'r', encoding='utf-8') as f:
            for line in f:
                line = line.strip()
                if line and not line.startswith('#'):  # Игнорируем пустые строки и комментарии
                    ignore_patterns.append(line)
        LOGGER.info("Загружено %s правил игнорирования из %s", len(ignore_patterns), ignore_file_path)
    except Exception as e:
        LOGGER.exception("Ошибка чтения файла %s: %s", ignore_file_path, e)

    return ignore_patterns


def should_ignore(path, ignore_patterns, is_directory=False):
    # Проверяет, нужно ли игнорировать файл или папку
    if not ignore_patterns:
        return False

    # Нормализуем путь для сравнения
    normalized_path = os.path.normpath(path)

    for pattern in ignore_patterns:
        # Если паттерн заканчивается на /, это относится только к директориям
        if pattern.endswith('/'):
            if not is_directory:
                continue
            pattern = pattern[:-1]  # Убираем trailing slash

        # Простое сравнение имен
        if pattern in normalized_path:
            return True

        # Поддержка wildcard *
        if '*' in pattern:
            import fnmatch
            if fnmatch.fnmatch(normalized_path, pattern):
                return True

    return False


def get_file_content(file_path):
    # Получает содержимое файла
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    except UnicodeDecodeError as e:
        LOGGER.warning("Файл %s не прочитан как UTF-8: %s. Пробуем latin-1", file_path, e)
        try:
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
        except Exception as e:
            LOGGER.exception("Ошибка чтения файла %s: %s", file_path, e)
            return f"Ошибка чтения файла {file_path}: {e}"
    except Exception as e:
        LOGGER.exception("Ошибка чтения файла %s: %s", file_path, e)
        return f"Ошибка чтения файла {file_path}: {e}"


def sanitize_text(text):
    # Удаляет из текста символы, несовместимые с XML
    if text is None:
        return ""
    return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', text)


def add_file_to_doc(doc, file_name, content, file_number, mode_settings):
    # Добавляет содержимое файла в документ.
    # Добавляем заголовок файла как Heading 1
    p = doc.add_paragraph()
    p.style = 'Heading 1'
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"ПРИЛОЖЕНИЕ {file_number}. {file_name}")
    apply_run_style(
        run,
        mode_settings.get("file_name_font_name"),
        mode_settings.get("file_name_font_size"),
        mode_settings.get("file_name_color"),
        mode_settings.get("file_name_bold"),
    )

    # Добавляем содержимое с уменьшенным шрифтом
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(mode_settings.get("content_space_after", 0))
    p.paragraph_format.line_spacing = mode_settings.get("content_line_spacing", 1)
    run = p.add_run(content)
    apply_run_style(
        run,
        mode_settings.get("content_font_name"),
        mode_settings.get("content_font_size"),
        mode_settings.get("content_color"),
    )

    # Добавляем разделитель
    doc.add_paragraph()


def log_walk_error(error):
    path = getattr(error, 'filename', '<неизвестно>')
    LOGGER.error("Ошибка обхода директории %s: %s", path, error)


def add_file_info_to_doc(directory, doc, use_two_columns=False, ignore_patterns=None, settings=None):
    # Рекурсивно проходит по директориям, получает имена и содержимое файлов и добавляет их в документ
    if ignore_patterns is None:
        ignore_patterns = []
    if settings is None:
        settings = DEFAULT_SETTINGS

    all_files = []

    # Сначала собираем все файлы с учетом игнорирования
    for root, dirs, files in os.walk(directory, onerror=log_walk_error):
        # Фильтруем директории для игнорирования
        try:
            dirs[:] = [d for d in dirs if not should_ignore(os.path.join(root, d), ignore_patterns, True)]
        except Exception as e:
            LOGGER.exception("Ошибка фильтрации директорий в %s: %s", root, e)
            raise

        for file in files:
            file_path = os.path.join(root, file)

            try:
                # Проверяем, нужно ли игнорировать файл
                if should_ignore(file_path, ignore_patterns, False):
                    LOGGER.info("Игнорируем: %s", file_path)
                    continue

                file_content = get_file_content(file_path)
                sanitized_content = sanitize_text(file_content)
                all_files.append((file_path, file, sanitized_content))
            except Exception as e:
                LOGGER.exception("Ошибка обработки файла %s: %s", file_path, e)
                all_files.append((file_path, file, f"Ошибка обработки файла {file_path}: {e}"))

    if not all_files:
        doc.add_paragraph("Не найдено файлов для обработки (возможно, все игнорируются)")
        return

    mode_settings = settings["two_columns"] if use_two_columns else settings["one_column"]

    for file_number, (file_path, file_name, content) in enumerate(all_files, start=1):
        try:
            add_file_to_doc(doc, file_name, content, file_number, mode_settings)
        except Exception as e:
            LOGGER.exception("Ошибка добавления файла %s в документ: %s", file_path, e)
            add_file_to_doc(
                doc,
                file_name,
                f"Ошибка добавления файла {file_path} в документ: {e}",
                file_number,
                mode_settings,
            )


def main(target_directory, output_file, use_two_columns=False):
    configure_logging()
    LOGGER.info("Начинаем обработку директории %s", target_directory)
    settings = load_settings()

    # Главная функция, создает документ, добавляет в него информацию о файлах и сохраняет его
    # Загружаем правила игнорирования
    try:
        ignore_patterns = load_ignore_patterns()
    except FileNotFoundError as e:
        LOGGER.warning("Внимание: %s", e)
        LOGGER.warning("Продолжаем без игнорирования файлов")
        ignore_patterns = []

    try:
        doc = docx.Document()
        document_settings = settings["document"]
        layout_settings = settings["layout"]

        # Добавляем заголовок с информацией о настройках
        title = doc.add_paragraph()
        title_run = title.add_run(document_settings.get("title", "Листинг проекта"))
        apply_run_style(
            title_run,
            document_settings.get("title_font_name"),
            document_settings.get("title_font_size"),
            document_settings.get("title_color"),
            document_settings.get("title_bold"),
        )

        info = doc.add_paragraph()
        info_font_name = document_settings.get("info_font_name")
        info_font_size = document_settings.get("info_font_size")
        info_color = document_settings.get("info_color")
        for text in (
            f"Директория: {target_directory}\n",
            f"Колонки: {'2' if use_two_columns else '1'}\n",
            f"Правил игнорирования: {len(ignore_patterns)}\n",
        ):
            info_run = info.add_run(text)
            apply_run_style(info_run, info_font_name, info_font_size, info_color)
        doc.add_paragraph()

        if use_two_columns:
            content_section = doc.add_section(WD_SECTION.CONTINUOUS)
            set_section_columns(
                content_section,
                2,
                layout_settings.get("two_columns_spacing_inches", 0.25),
            )

        add_file_info_to_doc(target_directory, doc, use_two_columns, ignore_patterns, settings)
        doc.save(output_file)
        LOGGER.info("Информация о файлах сохранена в '%s'", output_file)
        if use_two_columns:
            LOGGER.info("Режим: две колонки")
        else:
            LOGGER.info("Режим: одна колонка")
    except Exception as e:
        LOGGER.exception(
            "Обработка прервана. Директория: %s, выходной файл: %s. Причина: %s",
            target_directory,
            output_file,
            e,
        )
        raise


def ask_yes_no_question(question):
    # Задает вопрос да/нет и возвращает булево значение
    while True:
        answer = input(f"{question} (да/нет): ").lower().strip()
        if answer in ['да', 'д', 'yes', 'y']:
            return True
        elif answer in ['нет', 'н', 'no', 'n']:
            return False
        else:
            print("Пожалуйста, введите 'да' или 'нет'")


if __name__ == "__main__":
    target_directory = input("Введите путь к директории: ")  # Запрашиваем путь к директории у пользователя
    output_file = input("Введите имя выходного файла Word (.docx): ")  # Запрашиваем имя выходного файла

    # Спрашиваем о двухколоночном режиме
    use_two_columns = ask_yes_no_question("Использовать двухколоночный режим для экономии места?")

    main(target_directory, output_file, use_two_columns)
